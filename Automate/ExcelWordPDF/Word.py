import docx, os
os.chdir(r"C:\Users\DavidLapadula\Downloads")

d = docx.Document("demo.docx")
p = d.paragraphs[1]

p.runs[3].underline = True
p.style = "Title"

d2 = docx.Document()
d2.add_paragraph("This is a paragraph")
d2.paragraph[0].add_run("New run")

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return "\n".join(fullText)

print(getText("demo.docx"))



